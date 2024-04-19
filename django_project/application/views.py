import os
import re
from docx import Document
from call.models import Call
from rest_framework import status
from .permissions import IsStudent
from django.http import JsonResponse
from rest_framework import permissions
from rest_framework.response import Response
from django_project.constants import Constants
from rest_framework.decorators import api_view, permission_classes

from student.views import ApplicationDataView


@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated, IsStudent])
def get_region_call(request):
    call_id = request.query_params.get('call')
    call = Call.objects.get(id=call_id)
    region = call.university_id.get_region_display()
    if region == 'Uniandes':
        return JsonResponse({'region': 'Uniandes'}, status=status.HTTP_200_OK)
    elif region == 'Convenio Sigueme/Nacional':
        return JsonResponse({'region': 'Nacional'}, status=status.HTTP_200_OK)
    return JsonResponse({'region': 'Internacional'}, status=status.HTTP_200_OK)


@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated, IsStudent])
def create_forms(request):

    # Set the contact person and health information of the student
    response = ApplicationDataView().put(request)
    if response.status_code != status.HTTP_200_OK:
        return Response(response.data, status=status.HTTP_400_BAD_REQUEST)

    # Check student has a contact_id
    student = request.user.student
    contact_person = student.contact_id
    if contact_person is None:
        return JsonResponse({'error': 'La persona de contacto no ha sido definida'}, status=status.HTTP_400_BAD_REQUEST)

    # Info student
    attributes = dict()
    for attribute in student._meta.get_fields():
        try:
            # TODO: add contact_info to list
            if 'contact' in attribute.name:
                continue

            elif attribute.choices is not None:  # It is an enum
                display_method_name = f'get_{attribute.name}_display'
                attributes[attribute.name.upper()] = getattr(student, display_method_name)()
            else:
                value = getattr(student, attribute.name)
                if value is None:  # The value has not been set (like diseases)
                    value = 'No hay'
                attributes[attribute.name.upper()] = value
        except AttributeError:  # It is like a many-to-one relation
            continue

    # Info coordinator
    headquarter = student.get_headquarter_display()
    faculty = student.get_faculty_display()
    major = student.get_major_display()
    info_coordinator = Constants.INFO_FACULTIES[headquarter][faculty][major]

    path_original_forms = os.path.join('forms', 'original_forms')
    path_filled_forms = os.path.join('forms', 'filled_forms')
    for name_form in os.listdir(path_original_forms):
        path_doc = os.path.join(path_original_forms, name_form)
        doc = Document(path_doc)

        # TODO: fix keep the original format
        # TODO: make it more efficient
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text in attributes:
                        pattern = r'\b{}\b'.format(re.escape(cell.text))
                        cell.text = re.sub(pattern, f'{attributes[cell.text]}', cell.text)

        for paragraph in doc.paragraphs:
            for word in paragraph.text.split():
                if word in attributes:
                    pattern = r'\b{}\b'.format(re.escape(word))
                    paragraph.text = re.sub(pattern, f'{attributes[word]}', paragraph.text)

        doc.save(os.path.join(path_filled_forms, name_form))

    return JsonResponse({'result': 'Formulario regio'}, status=status.HTTP_200_OK)
