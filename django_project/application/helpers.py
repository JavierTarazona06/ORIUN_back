import os
import re
import pytz
import shutil
from math import ceil
from _io import BytesIO
from typing import Union
from docx import Document
from call.models import Call
from datetime import datetime
from datetime import timedelta
from urllib.parse import quote
from google.cloud import storage
from django.conf import settings
from django.core.exceptions import ImproperlyConfigured
from student.models import Student
from rest_framework.request import Request
from application.models import Application
from data.constants import Constants
import logging

logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

_STORAGE_CLIENT = None


def _storage_backend() -> str:
    return getattr(settings, "STORAGE_BACKEND", "local")


def _get_storage_client():
    global _STORAGE_CLIENT
    if _STORAGE_CLIENT is not None:
        return _STORAGE_CLIENT

    credentials = getattr(settings, "GOOGLE_APPLICATION_CREDENTIALS", "")
    if not credentials:
        raise ImproperlyConfigured(
            "GOOGLE_APPLICATION_CREDENTIALS is required when STORAGE_BACKEND=gcs."
        )
    _STORAGE_CLIENT = storage.Client.from_service_account_json(credentials)
    return _STORAGE_CLIENT


def _local_bucket_path(name_bucket: str) -> str:
    path = os.path.join(settings.LOCAL_STORAGE_ROOT, name_bucket)
    os.makedirs(path, exist_ok=True)
    return path


def _find_local_file(name_bucket: str, file_name: str) -> str:
    bucket_path = _local_bucket_path(name_bucket)
    for name in sorted(os.listdir(bucket_path)):
        if name.startswith(file_name):
            return os.path.join(bucket_path, name)
    raise FileNotFoundError(f"No file was found with the requested name: {file_name}")


def _publish_local_original_doc(file_name: str) -> str:
    source_dirs = [
        os.path.join(settings.BASE_DIR, "data", "forms", "templates"),
        os.path.join(settings.BASE_DIR, "data", "docs_applications", "general"),
    ]
    destination_dir = _local_bucket_path("original_doc")

    for source_dir in source_dirs:
        if not os.path.isdir(source_dir):
            continue
        for name in sorted(os.listdir(source_dir)):
            if name.startswith(file_name):
                source = os.path.join(source_dir, name)
                destination = os.path.join(destination_dir, name)
                shutil.copyfile(source, destination)
                return destination

    raise FileNotFoundError(f"No original document was found with the requested name: {file_name}")


def set_variables(request: Request, student: Student) -> dict[str, Union[str, list]]:
    """
    Returns a dictionary where the keys are the uppercase strings that should be replaced on the
    forms and the values are the strings that will be used to replace those upper strings. It includes
    the personal information of the student, info of its coordinator, info about the destination
    university and/or information about the courses (this one is a list).
    """
    # Info student
    attributes = dict()
    for attribute in student._meta.get_fields():
        try:
            if attribute.name == 'user':  # It is the user info
                user = getattr(student, attribute.name)
                attributes['FIRST_NAME'] = getattr(user, 'first_name')
                attributes['LAST_NAME'] = getattr(user, 'last_name')
                attributes['EMAIL'] = getattr(user, 'email')

            if 'contact' in attribute.name:  # It is the contact person info
                contact_person = getattr(student, attribute.name)
                for attr_contact in contact_person._meta.get_fields():
                    try:
                        name_attr = attr_contact.name.upper()
                        attributes[f'CONTACT_{name_attr}'] = getattr(contact_person, attr_contact.name)
                    except AttributeError:  # It is like a many-to-one relation
                        continue
            elif attribute.choices is not None:  # It is an enum
                display_method_name = f'get_{attribute.name}_display'
                attributes[attribute.name.upper()] = getattr(student, display_method_name)()
            else:
                value = getattr(student, attribute.name)
                if value is None:  # The value has not been set (like diseases)
                    value = 'No hay'
                attributes[attribute.name.upper()] = value
        except AttributeError:  # It is like a many-to-one relation
            continue

    # Info coordinator
    headquarter = student.get_headquarter_display()
    faculty = student.get_faculty_display()
    major = student.get_major_display()
    try:
        info_coordinator = Constants.INFO_FACULTIES[headquarter][faculty][major]
        attributes['NAME_COORD_UNAL'] = info_coordinator['Coordinador Curricular']
        attributes['PHONE_COORD_UNAL'] = info_coordinator['Teléfono Coordinador']
        attributes['EMAIL_COORD_UNAL'] = info_coordinator['Correo Coordinador']
    except KeyError:
        attributes['NAME_COORD_UNAL'] = ''
        attributes['PHONE_COORD_UNAL'] = ''
        attributes['EMAIL_COORD_UNAL'] = ''

    # Info destination university
    call = Call.objects.get(id=request.data['call'])
    attributes['DEST_INSTITUTION'] = call.university.name
    attributes['DEST_COUNTRY'] = call.university.country
    attributes['DEST_CITY'] = call.university.city
    info_mobility = request.data.get('info_mobility')
    if info_mobility is not None:
        for name, info in info_mobility.items():
            if 'date' in name:
                info = datetime.strptime(info, '%d-%m-%Y')
            attributes[f'{name.upper()}'] = info
    if 'START_DATE' and 'END_DATE' in attributes:
        attributes['DURATION'] = ceil((attributes['END_DATE'] - attributes['START_DATE']).days / 30)
    if 'START_DATE' in attributes:
        attributes['START_DATE'] = attributes['START_DATE'].strftime('%d-%m-%Y')
    if 'END_DATE' in attributes:
        attributes['END_DATE'] = attributes['END_DATE'].strftime('%d-%m-%Y')

    # Info courses
    info_courses = request.data.get('info_courses')
    if info_courses is not None:
        attributes['INFO_COURSES'] = []
        for course in info_courses:
            info_course = {}
            for name, info in course.items():
                info_course[f'{name.upper()}'] = info
            attributes['INFO_COURSES'].append(info_course)

    time_location = pytz.timezone('America/Bogota')
    day, month, year = datetime.now(time_location).strftime('%d.%m.%Y').split('.')
    attributes['DATE'], attributes['MONTH'], attributes['YEAR'] = day, month, year

    attributes['BIRTH_DATE'] = str(attributes['BIRTH_DATE'])

    # Use only the first course info
    try:
        info_courses = attributes['INFO_COURSES'][0]
        attributes['CODE_UNAL'] = info_courses['CODE_UNAL']
        attributes['CODE_DESTINY'] = info_courses['CODE_DESTINY']
        attributes['NAME_UNAL'] = info_courses['NAME_UNAL']
        attributes['NAME_DESTINY'] = info_courses['NAME_DESTINY']
    except KeyError:
        attributes['CODE_UNAL'] = ''
        attributes['CODE_DESTINY'] = ''
        attributes['NAME_UNAL'] = ''
        attributes['NAME_DESTINY'] = ''

    return attributes


def fill_forms(attributes: dict[str, str], path_original_forms: str, path_save_forms: str) -> None:
    """
    Replaces the keys of the given dictionary with the values of that dictionary in each form and
    saves them to the given path.
    """
    if not os.path.exists(path_save_forms):
        os.mkdir(path_save_forms)

    for name_form in os.listdir(path_original_forms):
        path_doc = os.path.join(path_original_forms, name_form)
        doc = Document(path_doc)

        # TODO: make filling of courses dynamic (does not depend on # of courses)
        # TODO: fix keep the original format
        # TODO: make it more efficient
        # TODO: finish filling up the forms with the uppercase words
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text in attributes:
                        pattern = r'\b{}\b'.format(re.escape(cell.text))
                        cell.text = re.sub(pattern, f'{attributes[cell.text]}', cell.text)

        for paragraph in doc.paragraphs:
            for word in paragraph.text.split():
                if word in attributes:
                    pattern = r'\b{}\b'.format(re.escape(word))
                    paragraph.text = re.sub(pattern, f'{attributes[word]}', paragraph.text)

        name, extension = name_form.split('.')
        new_name = f'{name}_{os.path.basename(path_save_forms)}.{extension}'
        doc.save(os.path.join(path_save_forms, new_name))


def upload_forms(path_save_forms: str) -> None:
    """
    Uploads generated forms to the configured storage backend, then removes the temporary folder.
    """
    list_forms = []
    for form in os.listdir(path_save_forms):
        path_form = os.path.join(path_save_forms, form)
        if _storage_backend() == "local":
            destination = os.path.join(_local_bucket_path("filled_doc"), form)
            shutil.copyfile(path_form, destination)
        else:
            bucket = _get_storage_client().bucket('filled_doc')
            blob = bucket.blob(form)
            blob.upload_from_filename(path_form)
        list_forms.append(path_form)

    for form in list_forms:
        os.remove(form)
    os.rmdir(path_save_forms)


def get_link_file(type_file: str, file_name: str) -> str:
    """
    Return a public link to an object in a bucket that can be used from 10 minutes. If the file
    does not exist, a FileNotFoundError is raised. No need to send the extension of the file in
    its name.
    """
    if _storage_backend() == "local":
        try:
            path = _find_local_file(type_file, file_name)
        except FileNotFoundError:
            if type_file != "original_doc":
                raise
            path = _publish_local_original_doc(file_name)
        return f"{settings.LOCAL_STORAGE_URL_PREFIX}/{type_file}/{quote(os.path.basename(path))}"

    bucket = _get_storage_client().bucket(type_file)
    blobs = bucket.list_blobs(prefix=file_name)
    try:
        blob = next(blobs)
    except StopIteration:
        raise FileNotFoundError(f"No se encontró el archivo con el nombre especificado: {file_name}")

    # Set the expiration time for the signed URL
    url_expiration = timedelta(minutes=10)
    signed_url = blob.generate_signed_url(expiration=url_expiration, method='GET')

    return signed_url


def upload_object(name_bucket: str, source_file: BytesIO, destination_blob_name: str) -> None:
    """
    Uploads a file-like object to the configured storage backend.
    """
    if hasattr(source_file, "seek"):
        source_file.seek(0)

    if _storage_backend() == "local":
        path = os.path.join(_local_bucket_path(name_bucket), destination_blob_name)
        with open(path, "wb") as destination:
            shutil.copyfileobj(source_file, destination)
        return

    bucket = _get_storage_client().bucket(name_bucket)
    blob = bucket.blob(destination_blob_name)
    blob.upload_from_file(source_file)


def delete_object(name_bucket: str, name_file: str):
    """
    Deletes a document from a bucket.
    """
    if _storage_backend() == "local":
        try:
            os.remove(_find_local_file(name_bucket, name_file))
        except FileNotFoundError:
            return
        return

    bucket = _get_storage_client().bucket(name_bucket)
    blobs = bucket.list_blobs(prefix=name_file)
    try:
        blob = next(blobs)
        blob.delete()
    except StopIteration:
        return


def check_docs(student: Student, call: Call) -> tuple[bool, list[str]]:
    """
    Used to make sure all the documents (depending on the region of the call) have been uploaded.
    Returns True if the document has been uploaded, False otherwise. If the return is True then it
    also returns the name of the missing documents.
    """
    # Name of the documents that should be found on GCP
    name_documents = Application.name_docs

    # Check national or international documents, if needed
    region = call.university.get_region_display()
    print(region, name_documents, student)
    logger.info("Initial name_documents: %s %s", region, name_documents)
    logger.debug("Initial name_documents: %s %s", region, name_documents)
    logger.error("Initial name_documents: %s %s", region, name_documents)
    if region == 'Convenio Sigueme/Nacional':
        name_documents.extend(Application.national_name_docs)

    elif region != 'Uniandes':  # It is an international call
        name_documents.extend(Application.international_name_docs)

    # Check each possible file
    missing_docs = []
    for name_doc in name_documents:
        complete_name = f'{name_doc}_{student.id}_{call.id}'
        if not exists(complete_name):
            missing_docs.append(name_doc)

    if len(missing_docs) > 0:
        return False, missing_docs
    return True, missing_docs


def exists(file_name: str) -> bool:
    """
    Returns True if the given file exists in the complete_doc bucket, False otherwise.
    """
    if _storage_backend() == "local":
        try:
            _find_local_file("complete_doc", file_name)
            return True
        except FileNotFoundError:
            return False

    bucket = _get_storage_client().bucket('complete_doc')
    blobs = bucket.list_blobs(prefix=file_name)
    try:
        _ = next(blobs)
        return True
    except StopIteration:
        return False
